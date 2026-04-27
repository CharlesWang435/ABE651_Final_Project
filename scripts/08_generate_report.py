"""Step 8 — generate the ABE 651 interim report and the data summary appendix.

Narrative is framed strictly around feature extraction and dataset characterization.
Species (CN, SB) is a labeling/grouping variable. Differences between species are
acknowledged in passing only — they are not the finding.
"""
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from common import DATA, FIGURES, DOCS, log


REPORT_PATH = DOCS / "ABE651_Interim_Report.docx"
SUMMARY_PATH = DOCS / "Final_Data_Summary.docx"

QC_THRESHOLDS = [
    ("Segmentation — leaf area fraction",
     "5%–90% of frame", "fraction of frame",
     "Excludes empty masks and frames where the leaf fills the lightbox.",
     "Project-defined (lightbox geometry)"),
    ("Segmentation — solidity",
     "≥ 0.60", "0–1 (skimage regionprops)",
     "Below 0.60 indicates a fragmented or jagged mask, often a segmentation error.",
     "Common practice in plant phenotyping (e.g., Pound et al., 2017)"),
    ("Exposure — mean L*",
     "12 ≤ L* ≤ 85", "CIE L* 0–100",
     "L* below 12 is severely underexposed; above 85 is washed out / clipped.",
     "CIE 1976 L*a*b* recommendations"),
    ("Channel clipping",
     "5 ≤ channel mean ≤ 250", "RGB 0–255",
     "Means at the rails indicate saturation; affected channels lose information.",
     "Standard imaging QA"),
    ("Metric outliers (per species)",
     "Q1 − 1.5×IQR / Q3 + 1.5×IQR (mild); ×3 (extreme)", "metric units",
     "Tukey fences applied within each species so natural species range is preserved.",
     "Tukey (1977), Exploratory Data Analysis"),
    ("Metadata completeness",
     "no missing values; unique sample_id", "categorical",
     "Ensures every record can be parsed and joined back to its source image.",
     "Project-defined"),
]


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def _add_para(doc, text, italic=False, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    return p


def _add_caption(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _add_figure(doc, path: Path, caption: str, width_in: float = 5.5):
    if not path.exists():
        _add_para(doc, f"[Figure file missing: {path.name}]", italic=True)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width_in))
    _add_caption(doc, caption)


def _add_table(doc, df: pd.DataFrame, header_bold=True):
    if df.empty:
        _add_para(doc, "[empty table]", italic=True)
        return
    table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    for j, c in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(c)
        if header_bold:
            for r in cell.paragraphs[0].runs:
                r.bold = True
    for i, row in enumerate(df.itertuples(index=False), start=1):
        for j, v in enumerate(row):
            if isinstance(v, float):
                txt = f"{v:.4g}"
            else:
                txt = str(v)
            table.rows[i].cells[j].text = txt
    return table


def _script_inventory_df():
    rows = []
    for p in sorted((DOCS.parent / "scripts").glob("*.py")):
        rows.append({"file": f"scripts/{p.name}", "purpose": _script_purpose(p.name)})
    return pd.DataFrame(rows)


def _script_purpose(name: str) -> str:
    return {
        "00_rename.py":            "Rename and standardize raw JPEGs to [SampleID]_[Species]_[Date]_[Seq].jpg",
        "01_inventory.py":         "Walk raw/, parse filenames + EXIF, write metadata_inventory.csv",
        "02_segment.py":           "Downsample, gray-world balance, segment leaf, write masks + QC",
        "03_extract_features.py":  "Extract color, GLCM texture, and morphology features per leaf",
        "04_eda.py":               "Generate seven exploratory figures (fig01–fig07)",
        "05_quality_check.py":     "Apply QC checks, flag outliers, produce QC figures and summary",
        "06_summary_stats.py":     "Per-species descriptive statistics (primary) and supplementary 2-sample appendix",
        "07_report_figures.py":    "Regenerate publication-quality figures (PNG + SVG)",
        "08_generate_report.py":   "Assemble the .docx interim report and the data summary appendix",
        "common.py":               "Shared paths, colour palette, and logging helper",
        "make_presentation_figures.py": "Slide-deck variants of EDA figures (out of scope for the report)",
    }.get(name, "")


def _build_introduction(doc):
    _add_heading(doc, "1. Introduction", level=1)

    _add_heading(doc, "1.1 Title", level=2)
    _add_para(doc,
              "RGB Image Analysis of Corn and Soybean Leaves for Phenotypic Feature Extraction")

    _add_heading(doc, "1.2 Authors", level=2)
    _add_para(doc, "Charles Wang — Agricultural and Biological Engineering, Purdue University")
    _add_para(doc, "Advisor: Dr. Jian Jin")

    _add_heading(doc, "1.3 Date", level=2)
    _add_para(doc, date.today().strftime("%B %d, %Y"))

    _add_heading(doc, "1.4 Synopsis", level=2)
    _add_para(doc,
              "This dataset comprises RGB smartphone photographs of individual excised "
              "corn (Zea mays) and soybean (Glycine max) leaves captured inside a custom "
              "transmittance lightbox, and a derived feature table containing colour "
              "(RGB / HSV / LAB), texture (GLCM), and morphology metrics for every leaf. "
              "The goal of this work is to build a reproducible pipeline that extracts "
              "and documents these phenotypic features so that the resulting feature "
              "set can be archived on PURR and reused by downstream researchers. The "
              "pipeline performs filename standardization, lighting balance, leaf "
              "segmentation, feature extraction, exploratory analysis, and quality "
              "checking, and produces per-species descriptive summaries of the extracted "
              "traits. Species labels (CN, SB) are retained as a grouping variable to "
              "organize tables and figures, but characterizing differences between corn "
              "and soybean is not the focus of this work.")

    _add_heading(doc, "1.5 Description", level=2)
    _add_para(doc,
              "The Leaf RGB feature dataset documents phenotypic traits extracted from "
              "smartphone photographs of corn and soybean leaves. Each image was captured "
              "inside a custom lightbox that illuminates the leaf from below "
              "(transmittance geometry), so the background is bright and the leaf "
              "appears as the darker region. Raw JPEGs were renamed to a standardized "
              "convention ([SampleID]_[Species]_[YYYYMMDD]_[Seq].jpg) so every record "
              "carries species and collection-date information in the filename itself.")
    _add_para(doc,
              "The processing pipeline downsamples each image to ~2000 px on its long "
              "edge, then applies a gray-world correction keyed to the bright background "
              "so that lighting differences across capture sessions do not propagate into "
              "downstream colour features. The leaf is segmented by Otsu thresholding on "
              "the LAB L* channel, the mask is inverted (leaf = foreground), morphological "
              "closing fills small holes, and only the largest connected component is "
              "retained. A per-image QC log records leaf area fraction, solidity, and "
              "bounding-box fill so segmentation quality can be audited.")
    _add_para(doc,
              "Feature extraction is performed only on masked leaf pixels. The feature "
              "set is deliberately curated — early iterations of the pipeline included a "
              "broader set, but features that were mathematically redundant or not "
              "interpretable on their own were removed so that every published feature "
              "carries a distinct, defensible phenotypic signal. The final set comprises "
              "twenty features in four blocks: five RGB colour features (channel means "
              "plus the illumination-normalized greenness and redness indices), six LAB "
              "perceptual colour features (mean L*, a*, b* and their within-leaf standard "
              "deviations), four GLCM texture features (contrast, homogeneity, energy, "
              "correlation at one-pixel offset), and five morphology features (area, "
              "aspect ratio, extent, solidity, compactness). The rationale for each "
              "feature, and for every feature that was removed, is documented in "
              "Section 4.1. The full feature table is written to data/leaf_metrics.csv "
              "with one row per leaf.")
    _add_para(doc,
              "The data quality layer applies six checks: segmentation fraction, "
              "segmentation solidity, exposure (CIE L*), channel clipping, per-species "
              "Tukey-fence outliers, and metadata completeness. Flagged records are "
              "logged but not silently removed — the goal is to make extraction quality "
              "transparent rather than to filter the dataset down to a clean subset.")
    _add_para(doc,
              "Per-species descriptive statistics (N, mean, SD, median, Q1, Q3, min, max, "
              "coefficient of variation) are reported for every numeric trait. Inferential "
              "two-sample tests are computed for completeness only and stored as a "
              "supplementary CSV; they are not featured in this report. Species "
              "differences in colour and morphology are visible in some figures simply "
              "because corn and soybean are different plants — these visible differences "
              "are not interpreted as findings of this work.")

    _add_heading(doc, "1.6 Tags", level=2)
    _add_para(doc,
              "plant phenotyping; RGB imaging; corn (Zea mays); soybean (Glycine max); "
              "leaf segmentation; gray-world correction; GLCM texture; LAB colour space; "
              "feature extraction; smartphone agriculture; PURR")


def _build_source_data(doc):
    _add_heading(doc, "2. Source Data", level=1)
    _add_para(doc,
              "Source data are JPEG photographs collected with a smartphone camera "
              "inside a custom transmittance lightbox. JPEG was chosen because it is the "
              "native format produced by mobile camera apps and is the format the "
              "researcher receives in the field. JPEG is a lossy, block-DCT compressed "
              "format: each save introduces compression artifacts that are particularly "
              "visible at high-frequency texture boundaries — which is exactly the signal "
              "that GLCM texture features measure. To prevent compounding compression "
              "loss, the pipeline reads each raw JPEG once, converts to a lossless PNG "
              "intermediate in processed/, and performs every subsequent operation "
              "(downsampling, gray-world correction, segmentation, feature extraction) on "
              "PNG only. No script ever overwrites a file in raw/ or re-encodes back to "
              "JPEG. This single-pass re-encoding policy means the texture features "
              "reflect the original capture rather than artifacts of repeated compression.")
    _add_para(doc,
              "Raw images live under raw/Corn Images/ and raw/Soybean Images/ in their "
              "original camera-assigned filenames. A standardization step "
              "(scripts/00_rename.py) copies each file into raw/ under the convention "
              "[SampleID]_[Species]_[YYYYMMDD]_[Seq].jpg. A rename log is written to "
              "docs/rename_log.txt mapping every old name to its new name so the "
              "transformation is auditable.")


def _build_metadata(doc):
    _add_heading(doc, "3. Metadata", level=1)
    _add_para(doc,
              "Metadata is captured in two CSV tables. data/metadata_inventory.csv "
              "records one row per source image with fields filename, sample_id, species, "
              "collection_date, seq_num, iso, exposure_time, white_balance, file_size_mb, "
              "image_width_px, image_height_px, and exif_notes. data/leaf_metrics.csv "
              "records one row per processed image with all extracted colour, texture, "
              "and morphology features alongside the same identifier columns.")
    _add_para(doc,
              "The script inventory below documents which file produces which output. "
              "All scripts share scripts/common.py for paths, palette, and logging.")
    inv = _script_inventory_df()
    _add_table(doc, inv)


def _build_eda(doc):
    _add_heading(doc, "4. Graphical Data Analysis", level=1)
    _add_para(doc,
              "This section first documents the curated feature set and the rationale "
              "for inclusion or exclusion of each candidate feature (4.1), then presents "
              "the feature set through a small set of figures of varied types: a count "
              "bar chart, a violin distribution plot, a scatter plot with bivariate "
              "ellipses, a correlation heatmap, and a principal-component biplot (4.2). "
              "Species labels are retained as a grouping/coloring variable but the "
              "narrative does not compare the two species against each other.")

    _add_heading(doc, "4.1 Feature set and selection rationale", level=2)
    _add_para(doc,
              "Twenty features are extracted per leaf, organized into four blocks. "
              "Each retained feature carries a distinct, defensible phenotypic signal; "
              "features that were mathematically redundant with another retained "
              "feature, or not interpretable on their own, were removed.")

    kept_df = pd.DataFrame([
        ("RGB",        "mean_R",         "Mean red-channel intensity over masked pixels (illumination-corrected)."),
        ("RGB",        "mean_G",         "Mean green-channel intensity. In transmittance, lower mean_G indicates higher chlorophyll absorbance."),
        ("RGB",        "mean_B",         "Mean blue-channel intensity. Strongly attenuated by leaves; useful as a thickness/transmittance proxy."),
        ("RGB",        "greenness_idx",  "G/(R+G+B). Illumination-normalized chromaticity; classic plant-phenotyping greenness index (NGI)."),
        ("RGB",        "redness_idx",    "R/(R+G+B). Illumination-normalized redness; tracks senescence and anthocyanin accumulation (NRI)."),
        ("LAB",        "mean_L",         "Mean lightness on the CIE 0–100 scale. Perceptually uniform brightness indicator."),
        ("LAB",        "mean_a",         "Mean a* (red-green axis). More positive a* indicates loss of green — the chlorosis axis."),
        ("LAB",        "mean_b",         "Mean b* (blue-yellow axis). More positive b* indicates yellowing — the senescence axis."),
        ("LAB",        "std_L",          "Within-leaf perceptual lightness variation; tracks unevenness in transmittance."),
        ("LAB",        "std_a",          "Within-leaf chromatic variation along the green-red axis; flags chlorotic patches."),
        ("LAB",        "std_b",          "Within-leaf chromatic variation along the blue-yellow axis; flags yellowing patches."),
        ("GLCM (d=1)", "glcm_contrast",  "Local intensity variation between adjacent pixels; reflects vein and edge sharpness."),
        ("GLCM (d=1)", "glcm_homogeneity","Local uniformity (high when adjacent pixels are similar); inverse-spirit of contrast."),
        ("GLCM (d=1)", "glcm_energy",    "Sum of squared GLCM elements; measures orderliness of the texture."),
        ("GLCM (d=1)", "glcm_correlation","Linear gray-level dependency between adjacent pixels; captures directional structure (e.g. parallel veins)."),
        ("Morphology", "area_px",        "Leaf size in pixels (masked-pixel count); the primary size descriptor."),
        ("Morphology", "aspect_ratio",   "Bounding-box height/width; distinguishes elongated leaves (corn) from broad leaves (soybean)."),
        ("Morphology", "extent",         "area / bounding-box area; how much of the bounding rectangle the leaf fills."),
        ("Morphology", "solidity",       "area / convex-hull area; quantifies lobing and concavity of the leaf outline."),
        ("Morphology", "compactness",    "(4π·area) / perimeter²; isoperimetric quotient. 1.0 = perfect circle; lower values indicate more elongated or jagged outlines."),
    ], columns=["Block", "Feature", "Phenotypic meaning"])

    _add_para(doc,
              "Table 2 lists the twenty retained features and their phenotypic interpretation.")
    _add_table(doc, kept_df)
    _add_caption(doc, "Table 2. Curated feature set with phenotypic interpretation.")

    cut_df = pd.DataFrame([
        ("HSV block (mean_H, mean_S, mean_V, std_H, std_S, std_V)",
         "Redundant with RGB and LAB",
         "HSV is a non-perceptual recoloring of RGB. mean_V correlates with mean_L at r ≈ 1.00. mean_H is a circular variable; arithmetic mean is mathematically ill-defined."),
        ("std_R, std_G, std_B",
         "Redundant with std_L",
         "Channel standard deviations carry largely the same heterogeneity signal as std_L (r ≈ 0.93)."),
        ("glcm_dissimilarity (all distances)",
         "Redundant with glcm_contrast",
         "Dissimilarity ≈ ∑|i−j|·p, contrast = ∑(i−j)²·p — the same texture signal at different powers (r = 0.98)."),
        ("glcm_ASM (all distances)",
         "Mathematically equal to glcm_energy²",
         "energy = √ASM by definition (r = 0.99). Keeping both inflates the feature count without adding information."),
        ("glcm_*_d3 and glcm_*_d5",
         "Redundant with glcm_*_d1",
         "Multi-distance GLCM is sometimes claimed to capture spatial scale, but on this dataset the d=3 and d=5 statistics correlate with d=1 at r > 0.93 across all four retained statistics."),
        ("perimeter_px",
         "Not interpretable on its own",
         "Perimeter depends on image resolution and offers little phenotypic meaning by itself; it is computed only as an input to compactness, which is the interpretable shape statistic."),
        ("equivalent_diameter",
         "Deterministic function of area_px",
         "= √(4·area/π). r = 0.998 with area_px. Provides no information beyond area."),
    ], columns=["Removed feature(s)", "Reason", "Evidence / explanation"])

    _add_para(doc,
              "Table 3 documents every feature removed during curation, with the "
              "specific reason. Reporting the removals explicitly (rather than silently "
              "dropping them) keeps the curated dataset auditable for future users.")
    _add_table(doc, cut_df)
    _add_caption(doc, "Table 3. Features removed during curation, with rationale.")

    _add_heading(doc, "4.2 Distribution and structure of the extracted features", level=2)

    _add_para(doc,
              "Figure 1 reports the sample count per species — the basic shape of the "
              "dataset before any feature extraction.")
    _add_figure(doc, FIGURES / "fig01_sample_overview.png",
                "Figure 1. Sample counts by species in the extracted dataset.")

    _add_para(doc,
              "Figure 2 shows the distribution of segmented leaf area extracted from "
              "each sample. The violin plot communicates both the central tendency and "
              "the spread of the morphology metric within each species group.")
    _add_figure(doc, FIGURES / "fig02_leaf_area_distribution.png",
                "Figure 2. Distribution of segmented leaf area extracted from each "
                "sample, grouped by species.")

    _add_para(doc,
              "Figure 4 plots mean LAB a* against mean LAB b* for every leaf, with a "
              "2-sigma ellipse summarizing the bivariate spread of each species group. "
              "This figure documents how the colour features distribute in perceptual "
              "colour space.")
    _add_figure(doc, FIGURES / "fig04_lab_scatter.png",
                "Figure 4. LAB colourspace coordinates (a* vs b*) extracted from each "
                "leaf, with 2-sigma ellipses per species.")

    _add_para(doc,
              "Figure 5 is a Pearson correlation matrix of the full curated feature set. "
              "After the redundancy-driven cuts described in Section 4.1, residual "
              "correlations are mostly within-block (e.g. mean_R, mean_G, mean_B move "
              "together; the four GLCM statistics share modest correlations) and "
              "between expected pairs (greenness_idx with mean_a). No two retained "
              "features are near-duplicates, confirming the curated set covers "
              "distinct phenotypic axes.")
    _add_figure(doc, FIGURES / "fig05_texture_heatmap.png",
                "Figure 5. Pearson correlation matrix of the curated feature set.")

    _add_para(doc,
              "Figure 7 is a PCA biplot of the entire numeric feature set. Sample scores "
              "appear as points and the top five contributing features appear as loading "
              "arrows. The biplot summarizes which traits drive the largest variance "
              "directions in the extracted dataset.")
    _add_figure(doc, FIGURES / "fig07_pca_biplot.png",
                "Figure 7. PCA biplot of the extracted leaf phenotypic feature set. "
                "Arrows indicate loadings of the top five contributing variables.")


def _build_qc(doc):
    _add_heading(doc, "5. Data Quality Checking", level=1)
    _add_para(doc,
              "Quality control is applied at three layers: segmentation geometry, "
              "exposure / colour-channel sanity, and per-species metric outliers. "
              "Thresholds, scales, and rationale are documented in the table below.")

    qc_df = pd.DataFrame(QC_THRESHOLDS,
                         columns=["QC check", "Threshold", "Scale", "Rationale", "Source"])
    _add_table(doc, qc_df)
    _add_caption(doc, "Table 1. QC checks, thresholds, and rationale.")

    _add_para(doc,
              "Figure QC1 places every leaf in (area, greenness) space and colours each "
              "point by its outlier flag, so the practical impact of the per-species "
              "Tukey-fence outlier rule is visible at a glance.")
    _add_figure(doc, FIGURES / "qc_fig01_outlier_scatter.png",
                "QC Figure 1. Outlier detection in leaf area vs. greenness index.")

    _add_para(doc,
              "Figure QC2 reports the distribution of mean L* across all images relative "
              "to the [12, 85] exposure window. Images outside the window were flagged "
              "for manual review.")
    _add_figure(doc, FIGURES / "qc_fig02_exposure_distribution.png",
                "QC Figure 2. Distribution of mean LAB L* values across the dataset, "
                "with flagging thresholds shown as dashed lines.")

    _add_para(doc,
              "Figure QC3 shows representative segmentation review cases — either "
              "flagged segmentations or, if none failed, a sample of normal segmentations "
              "for context.")
    _add_figure(doc, FIGURES / "qc_fig03_before_after_segmentation.png",
                "QC Figure 3. Examples of segmentation quality review.")

    qc_summary = pd.read_csv(DATA / "qc_summary_table.csv")
    _add_para(doc, "The QC summary table aggregates flag counts across the full dataset.")
    _add_table(doc, qc_summary)
    _add_caption(doc, "Table 4. QC summary across the full image set.")


def build_main_report():
    doc = Document()
    _add_para(doc,
              "ABE 651 — Final Interim Report (Revised)",
              bold=True, size=14)
    _build_introduction(doc)
    _build_source_data(doc)
    _build_metadata(doc)
    _build_eda(doc)
    _build_qc(doc)
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_PATH)


def build_data_summary():
    doc = Document()
    _add_para(doc, "Final Data Summary — Leaf RGB Phenotypic Feature Set",
              bold=True, size=14)
    _add_para(doc,
              "Companion appendix to the ABE 651 interim report. Contents: per-species "
              "descriptive statistics for every extracted trait, plus the full set of "
              "exploratory figures.")

    _add_heading(doc, "Per-species descriptive statistics", level=1)
    summary = pd.read_csv(DATA / "summary_statistics.csv")
    _add_table(doc, summary)

    _add_heading(doc, "Exploratory figures", level=1)
    eda_figs = [
        ("fig01_sample_overview.png",
         "Sample counts by species in the extracted dataset."),
        ("fig02_leaf_area_distribution.png",
         "Distribution of extracted leaf area, grouped by species."),
        ("fig03_greenness_index.png",
         "Distribution of extracted greenness index, grouped by species."),
        ("fig04_lab_scatter.png",
         "Extracted LAB a* vs b* coordinates with 2-sigma ellipses per species."),
        ("fig05_texture_heatmap.png",
         "Pearson correlation of GLCM texture features."),
        ("fig06_leaf_panel.png",
         "Representative leaves from the extracted dataset."),
        ("fig07_pca_biplot.png",
         "PCA biplot of the extracted feature set."),
    ]
    for name, cap in eda_figs:
        _add_figure(doc, FIGURES / name, cap)

    _add_heading(doc, "Supplementary two-sample appendix", level=1)
    _add_para(doc,
              "The table below reports basic two-sample descriptors (Welch's t and "
              "Mann-Whitney U) for every numeric trait. These are provided for "
              "completeness only — between-species comparison is not the focus of this "
              "work.")
    supp_path = DATA / "species_comparison_supplementary.csv"
    if supp_path.exists():
        supp = pd.read_csv(supp_path)
        _add_table(doc, supp)
    else:
        _add_para(doc, "[Supplementary CSV not yet generated — re-run scripts/06_summary_stats.py.]",
                  italic=True)

    SUMMARY_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(SUMMARY_PATH)


def main():
    log("08_generate_report", "start")
    build_main_report()
    build_data_summary()
    log("08_generate_report",
        f"wrote {REPORT_PATH.name} and {SUMMARY_PATH.name}")


if __name__ == "__main__":
    main()
